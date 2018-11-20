import os

import docx

# read the sections
dir_files = [f for f in os.listdir("to_merge\\")]
d = docx.Document()
for file in dir_files:
    doc = docx.Document(docx="to_merge\\" + file)

    article_title = doc.paragraphs[0]
    d.add_heading(article_title.text)
    for p in doc.paragraphs[1:]:
        d.add_paragraph(p.text, p.style)

    # output the sections
import time
ts = time.time()
import datetime
st = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H.%M.%S')
filename = os.path.abspath(os.curdir) + "\\" 'merged-' + st + '.docx'
print('Saving {}...'.format(filename))
d.save(filename)
