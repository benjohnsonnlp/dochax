import os
import re
from collections import OrderedDict

import docx

# read the sections
dir_files = [f for f in os.listdir("fixtures\\")]
doc_dict = OrderedDict()
for file in dir_files:
    doc = docx.Document(docx="fixtures\\" + file)

    article_title = None
    for p in doc.paragraphs:
        rx = '^[0-9]:[0-9][0-9]'
        prog = re.compile(rx)
        text = p.text.strip()
        if prog.search(text):
            article_title = text
            doc_dict[article_title] = [p]
        elif article_title:
            if text == '#MICPEL#':
                article_title = None
            else:
                doc_dict[article_title].append(p)

# output the sections

for title, paragraphs in doc_dict.items():
    d = docx.Document()
    d.add_heading(paragraphs[0].text)
    for p in paragraphs[1:]:
        if p.text.strip != '':
            d.add_paragraph(p.text, p.style)
    clean_name = title.replace(':', '.').replace(',', '').replace('(', '').replace(')', '').replace('/', '')
    filename = os.path.abspath(os.curdir) + "\\to_merge\\" + clean_name[:128] + '.docx'
    print('Saving {}...'.format(filename))
    d.save(filename)
